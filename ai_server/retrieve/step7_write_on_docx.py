from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from retrieve.step6_adapt_or_not import adapt_or_not
def create_json_to_docx(path_pdf, collection_name):
    context_queries, product_keys = adapt_or_not(path_pdf, collection_name)
    print("=== Bắt đầu tạo file Word ===")
    create_docx_file(context_queries, product_keys)
    print("=== Đã tạo file Word thành công ===")


def create_docx_file(context_queries, product_keys):
    doc = Document()
    doc.add_heading("BẢNG TUYÊN BỐ ĐÁP ỨNG VỀ KỸ THUẬT", level=1)

    # Tạo bảng 6 cột
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = [
        "Hạng mục số", "Tên hàng hoá",
        "Thông số kỹ thuật và các tiêu chuẩn của hàng hoá trong E-HSMT",
        "Thông số kỹ thuật và các tiêu chuẩn của hàng hoá chào trong E-HSDT",
        "Hồ sơ tham chiếu", "Tình đáp ứng của hàng hoá"
    ]

    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Ghi từng dòng
    for product, hang_hoa_dict in product_keys.items():
        for idx, (ten_hang_hoa, items) in enumerate(hang_hoa_dict.items(), start=1):
            ma_ids = items[:-2]  # Các ID
            dap_ung = items[-2]  # ví dụ: "đáp ứng"
            ho_so = items[-1]    # tài liệu tham chiếu

            # Tổng hợp thông số kỹ thuật
            eh_smt = ""
            eh_hsdt = ""

            for ma in ma_ids:
                if ma in context_queries:
                    eh_smt += f"- {context_queries[ma]['yeu_cau_ky_thuat_chi_tiet']}\n"
                    eh_hsdt += f"- {context_queries[ma].get('kha_nang_dap_ung',"")}\n"

            # Tạo dòng mới
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = ten_hang_hoa
            row[2].text = eh_smt.strip()
            row[3].text = eh_hsdt.strip()
            row[4].text = ho_so
            row[5].text = dap_ung

            for cell in row:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    # Lưu file
    doc.save("D:/study/LammaIndex/output/bang_tuyen_bo_dap_ung2.docx")

create_json_to_docx("D:/study/LammaIndex/documents/test.pdf", "thong_tin_san_pham")




