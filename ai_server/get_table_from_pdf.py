import camelot
import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement

def table_to_json(table_data, table_info):
    """Convert table data to JSON format"""
    if not table_data:
        return {}
    
    headers = table_data[0] if table_data else []
    rows = table_data[1:] if len(table_data) > 1 else []
    
    json_data = {
        "metadata": {
            "source_file": table_info["source_file"],
            "page": table_info["page"],
            "order": table_info["order"],
            "total_rows": len(rows),
            "total_columns": len(headers)
        },
        "headers": headers,
        "data": rows
    }
    
    return json_data

def get_continued_tables(tables, threshold=50):
    continued_tables = {}
    previous_table = False
    group_counter = 0

    page_height = 792  # PDF typical height minus margins

    for i, table in enumerate(tables):
        if previous_table and table.page == previous_table.page + 1 and len(table.cols) == len(previous_table.cols):
            previous_table_bottom = previous_table._bbox[1]
            current_table_top = table._bbox[3]

            if previous_table_bottom < (threshold / 100) * page_height and current_table_top > (1 - threshold / 100) * page_height:
                if (continued_tables.get(group_counter) is None):
                    continued_tables[group_counter] = [previous_table]
                continued_tables[group_counter].append(table)
            else:
                group_counter += 1
        else:
            group_counter += 1

        previous_table = table

    continued_tables = [value for value in continued_tables.values()]
    return continued_tables

def set_repeat_table_header(row):
    """Set repeat table row on every new page."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "true")
    trPr.append(tblHeader)

def save_table_to_word(table_json, output_path):
    """Save table data to Word document (only first 3 columns)"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Bảng từ file: {table_json["metadata"]["source_file"]}', 0)
    
    # Add metadata information
    metadata_para = doc.add_paragraph()
    metadata_para.add_run('Thông tin bảng:\n').bold = True
    metadata_para.add_run(f'• Trang: {table_json["metadata"]["page"]}\n')
    metadata_para.add_run(f'• Thứ tự: {table_json["metadata"]["order"]}\n')
    metadata_para.add_run(f'• Tổng số dòng: {table_json["metadata"]["total_rows"]}\n')
    metadata_para.add_run(f'• Tổng số cột gốc: {table_json["metadata"]["total_columns"]}\n')
    metadata_para.add_run(f'• Số cột được lưu: 3 (3 cột đầu tiên)\n')
    
    doc.add_paragraph()
    
    if table_json["headers"] and table_json["data"]:
        headers_first_3 = table_json["headers"][:3]
        num_cols = min(3, len(table_json["headers"]))
        num_rows = len(table_json["data"]) + 1  # +1 for header
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers_first_3):
            if i < len(header_cells):
                header_cells[i].text = str(header)
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        # ✅ Đặt header lặp lại
        set_repeat_table_header(table.rows[0])

        # Add data rows
        for row_idx, row_data in enumerate(table_json["data"]):
            if row_idx + 1 < len(table.rows):
                row_cells = table.rows[row_idx + 1].cells
                row_data_first_3 = row_data[:3]
                for col_idx, cell_data in enumerate(row_data_first_3):
                    if col_idx < len(row_cells):
                        row_cells[col_idx].text = str(cell_data)
    
    doc.save(output_path)
    print(f"Đã lưu bảng (3 cột đầu tiên) vào file Word: {output_path}")

def get_biggest_table_and_save_to_word(pdf_path, output_word_path=None, threshold=50):
    """Extract largest table from PDF and save to Word document"""
    tables = camelot.read_pdf(pdf_path, flavor='lattice', pages='all')
    continued_tables = get_continued_tables(tables, threshold)

    pdf_file_name = os.path.splitext(os.path.basename(pdf_path))[0]

    if output_word_path is None:
        output_word_path = f"{pdf_file_name}_largest_table.docx"

    processed = []
    all_table_jsons = []

    for i, table in enumerate(tables):
        if table in processed: 
            continue

        is_continued = any(table in sublist for sublist in continued_tables)
        all_table_data = list(table.data)

        if is_continued:
            group_index = next(index for index, sublist in enumerate(continued_tables) if table in sublist)
            for continued_table in continued_tables[group_index]:
                if continued_table == table or continued_table in processed: 
                    continue
                all_table_data.extend(continued_table.data[1:] if len(continued_table.data) > 1 else [])
                processed.append(continued_table)

        table_info = {
            "source_file": pdf_file_name,
            "page": table.parsing_report['page'],
            "order": table.parsing_report['order']
        }
        
        json_data = table_to_json(all_table_data, table_info)
        all_table_jsons.append(json_data)
        
        processed.append(table)

    if all_table_jsons:
        largest_table = max(all_table_jsons, key=lambda x: x.get('metadata', {}).get('total_rows', 0))
        save_table_to_word(largest_table, output_word_path)
        print("Thông tin bảng lớn nhất:")
        print(json.dumps(largest_table, ensure_ascii=False, indent=2))
        return largest_table
    else:
        print("Không tìm thấy bảng nào trong file PDF.")
        return None

if __name__ == "__main__":
    pdf_file_path = "D:\\study\\LammaIndex\\documents\\test1.pdf"
    output_word_file = "D:\\study\\LammaIndex\\output\\largest_table.docx"
    get_biggest_table_and_save_to_word(pdf_file_path, output_word_file)
